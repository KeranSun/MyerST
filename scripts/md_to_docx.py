"""Convert the manuscript/cover letter markdown to submission-grade .docx.

Handles: # headings, bold/italic inline, <sup>/<sub> runs, \* literal
asterisk, references, figure legends, and a Times New Roman layout.
"""

import re
import sys

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INLINE = re.compile(
    r"(<sup>.*?</sup>|<sub>.*?</sub>|\*\*.*?\*\*|\*[^*]+?\*|\\?\*)")


def add_runs(par, text):
    """Add runs to paragraph with sup/sub/bold/italic support."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok == r"\*":
            par.add_run("*")
        elif tok.startswith("<sup>"):
            r = par.add_run(tok[5:-6])
            r.font.superscript = True
        elif tok.startswith("<sub>"):
            r = par.add_run(tok[5:-6])
            r.font.subscript = True
        elif tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = par.add_run(tok[1:-1])
            r.italic = True
        else:
            par.add_run(tok)


def style_doc(doc):
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        for m in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
            setattr(section, m, Cm(2.5))


def convert(md_path, docx_path, title_size=16):
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    style_doc(doc)
    first_title = True
    in_refs = False
    for ln in lines:
        s = ln.rstrip()
        if not s:
            continue
        if s == "---":
            continue
        if s.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(s[2:])
            r.bold = True
            r.font.size = Pt(title_size)
            if first_title:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                first_title = False
        elif s.startswith("## "):
            in_refs = s[3:].strip() == "References"
            p = doc.add_paragraph()
            r = p.add_run(s[3:])
            r.bold = True
            r.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(12)
        elif s.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(s[4:])
            r.bold = True
            r.font.size = Pt(11.5)
            p.paragraph_format.space_before = Pt(10)
        elif s.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s[2:])
        elif re.match(r"^\d+\. ", s) and not in_refs:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\. ", "", s))
        else:
            p = doc.add_paragraph()
            add_runs(p, s)
            if in_refs:
                p.paragraph_format.left_indent = Cm(0.6)
                p.paragraph_format.first_line_indent = Cm(-0.6)
                p.paragraph_format.space_after = Pt(2)
            else:
                p.paragraph_format.first_line_indent = Cm(0.75)
                p.paragraph_format.space_after = Pt(4)
    doc.save(docx_path)
    print(f"saved {docx_path} ({len(doc.paragraphs)} paragraphs)")


if __name__ == "__main__":
    convert("paper/manuscript_v1.md", "outputs/manuscript_v1.docx")
    convert("paper/cover_letter.md", "outputs/cover_letter.docx",
            title_size=13)
